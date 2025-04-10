import fitz  # PyMuPDF
import docx
from io import BytesIO

def extract_text_from_pdf(file):
    """Extract text from PDF file (path or Streamlit UploadedFile)"""
    text = ""
    if hasattr(file, 'read'):  # Streamlit UploadedFile
        with BytesIO(file.read()) as pdf_bytes:
            with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf_document:
                for page in pdf_document:
                    text += page.get_text()
    else:  # File path
        with fitz.open(file) as pdf_document:
            for page in pdf_document:
                text += page.get_text()
    return text

def extract_text_from_docx(file):
    """Extract text from DOCX file (path or Streamlit UploadedFile)"""
    if hasattr(file, 'read'):  # Streamlit UploadedFile
        doc = docx.Document(BytesIO(file.read()))
    else:  # File path
        doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text(file):
    """Wrapper function to extract text based on file type"""
    if hasattr(file, 'name'):  # Streamlit UploadedFile
        filename = file.name.lower()
    else:  # File path
        filename = str(file).lower()
        
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(file)
    else:
        raise ValueError("Unsupported file type")